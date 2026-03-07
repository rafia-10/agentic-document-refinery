"""
FastTextExtractor — Cheapest Extraction Strategy
=================================================
Handles digital PDFs (pypdf), HTML (stdlib html.parser), plain text, and
Markdown with no external vision or layout engine.

Confidence heuristic
--------------------
``confidence = useful_chars / total_chars × (1.0 - junk_ratio)``

Spatial Provenance
------------------
Since FastText does not have a layout engine, it synthesizes approximate
line-level bounding boxes by dividing the page vertically based on the
number of lines detected.
"""

from __future__ import annotations

import html.parser
import io
import logging
import re
from pathlib import Path
from typing import Any

from src.models.schemas import (
    BoundingBox,
    DocumentProfile,
    ExtractedDocument,
    OriginType,
)
from src.strategies.base import BaseExtractor, ExtractionResult

logger = logging.getLogger(__name__)

_USEFUL_CHAR_RE = re.compile(r"[A-Za-z0-9.,;:!?'\"\-\(\)\[\]\{\}]")


class _HTMLStripper(html.parser.HTMLParser):
    """Minimal HTML parser that collects visible text."""

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_tags = {"script", "style", "head", "meta", "link"}
        self._current_tag = ""

    def handle_starttag(self, tag: str, attrs: Any) -> None:
        self._current_tag = tag.lower()

    def handle_endtag(self, tag: str) -> None:
        self._current_tag = ""

    def handle_data(self, data: str) -> None:
        if self._current_tag not in self._skip_tags:
            stripped = data.strip()
            if stripped:
                self._parts.append(stripped)

    def get_text(self) -> str:
        return "\n".join(self._parts)


class FastTextExtractor(BaseExtractor):
    """
    Strategy 1: Fast, cheap text extraction.
    """

    @property
    def name(self) -> str:
        return "fast_text"

    def extract(
        self,
        file_path: str | Path,
        profile: DocumentProfile,
    ) -> ExtractionResult:
        path = Path(file_path)
        raw_bytes = path.read_bytes()
        warnings: list[str] = []

        if profile.origin_type in (OriginType.NATIVE_DIGITAL,):
            pages, page_texts, warns = self._extract_pdf(raw_bytes)
            warnings.extend(warns)
        elif profile.origin_type == OriginType.HTML:
            page_texts = [self._extract_html(raw_bytes)]
            pages = 1
        elif profile.origin_type == OriginType.DOCX:
            page_texts = [self._extract_docx(raw_bytes)]
            pages = 1
        else:
            decoded = raw_bytes.decode("utf-8", errors="replace")
            page_texts = [decoded]
            pages = 1

        full_text = "\n\n".join(t for t in page_texts if t)
        
        # Spatial Provenance synthesis
        bounding_boxes = self._synthesize_bboxes(page_texts)
        
        # Enhanced confidence
        confidence = self._compute_enhanced_confidence(full_text)

        if profile.origin_type == OriginType.SCANNED_IMAGE:
            warnings.append(
                "FastTextExtractor: scanned PDF detected — confidence capped; "
                "escalating to Vision is recommended."
            )
            confidence = min(confidence, 0.40)

        doc = ExtractedDocument(
            document_id=profile.document_id,
            full_text=full_text,
            full_text_by_page=page_texts,
            page_count=max(pages, 1),
            overall_confidence=confidence,
            warnings=warnings,
            bounding_boxes=bounding_boxes,
        )
        return ExtractionResult(
            document=doc,
            confidence=confidence,
            strategy_name=self.name,
            warnings=warnings,
        )

    # ------------------------------------------------------------------
    # Spatial Synthesis
    # ------------------------------------------------------------------

    @staticmethod
    def _synthesize_bboxes(page_texts: list[str]) -> dict[str, BoundingBox]:
        """
        Produce approximate line-level bboxes.
        Assumes full page width (x0=0, x1=1) and equal line height.
        """
        boxes: dict[str, BoundingBox] = {}
        for p_idx, text in enumerate(page_texts):
            lines = [l for l in text.splitlines() if l.strip()]
            if not lines:
                continue
            line_height = 1.0 / len(lines)
            for l_idx, _ in enumerate(lines):
                box_id = f"p{p_idx:03d}-l{l_idx:03d}"
                y0 = l_idx * line_height
                y1 = (l_idx + 1) * line_height
                boxes[box_id] = BoundingBox(x0=0.0, y0=y0, x1=1.0, y1=y1)
        return boxes

    # ------------------------------------------------------------------
    # Enhanced Confidence
    # ------------------------------------------------------------------

    @staticmethod
    def _compute_enhanced_confidence(text: str) -> float:
        """
        Confidence based on:
        1. Character yield (useful vs total)
        2. Junk penalty (lines with too many non-alphas)
        """
        if not text:
            return 0.0
        
        # 1. Base yield
        char_total = len(text)
        char_useful = len(_USEFUL_CHAR_RE.findall(text))
        base_yield = char_useful / char_total if char_total > 0 else 0.0
        
        # 2. Junk detection
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        if not lines:
            return round(base_yield, 4)

        junk_count = 0
        for line in lines:
            # Check for high numeric density or very short nonsense lines
            nums = len(re.findall(r"\d", line))
            alphas = len(re.findall(r"[a-zA-Z]", line))
            if len(line) > 5 and nums > (alphas * 2): # mostly numbers
                junk_count += 1
            elif len(line) < 3 and not line[0].isalnum(): # tiny noise
                junk_count += 1
        
        junk_ratio = junk_count / len(lines)
        
        # Penalty scaling: 1.0 at 0% junk, 0.5 at 50% junk
        penalty = max(0.2, 1.0 - junk_ratio)
        
        return round(min(base_yield * penalty, 1.0), 4)

    # ------------------------------------------------------------------
    # Format-specific extractors
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_pdf(raw_bytes: bytes) -> tuple[int, list[str], list[str]]:
        warnings: list[str] = []
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
            page_texts: list[str] = []
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text() or ""
                    page_texts.append(text.strip())
                except Exception as exc:
                    warnings.append(f"Page {i + 1}: text extraction failed — {exc}")
                    page_texts.append("")
            return len(reader.pages), page_texts, warnings
        except ImportError:
            warnings.append("pypdf missing; using raw scanning.")
            raw_str = raw_bytes.decode("latin-1", errors="replace")
            text_blocks = re.findall(r"BT\s*(.*?)\s*ET", raw_str, re.DOTALL)
            return 1, [" ".join(text_blocks)[:8192]], warnings

    @staticmethod
    def _extract_html(raw_bytes: bytes) -> str:
        detected = chardet.detect(raw_bytes[:4096])
        enc = detected.get("encoding") or "utf-8"
        html_str = raw_bytes.decode(enc, errors="replace")
        stripper = _HTMLStripper()
        stripper.feed(html_str)
        return stripper.get_text()

    @staticmethod
    def _extract_docx(raw_bytes: bytes) -> str:
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            import zipfile
            try:
                with zipfile.ZipFile(io.BytesIO(raw_bytes)) as zf:
                    xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
                    texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, re.DOTALL)
                    return " ".join(texts)
            except: return ""
